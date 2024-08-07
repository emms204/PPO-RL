from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math
import datetime
import pandas as pd

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def build_revision_table(document):
	
	revision_table = document.add_table(rows=1, cols=5)
	hdr_cells = revision_table.rows[0].cells
	hdr_cells[0].text = 'Re'
	hdr_cells[1].text = 'Date'
	hdr_cells[2].text = 'Description'
	hdr_cells[3].text = 'By'
	hdr_cells[4].text = 'Appr'
	make_rows_bold(revision_table.rows[0])

	#add blank rows
	row_cells = revision_table.add_row().cells
	row_cells = revision_table.add_row().cells
	row_cells = revision_table.add_row().cells
	row_cells = revision_table.add_row().cells

	#fill in bottom row
	row_cells[0].text = "0"
	row_cells[1].text = datetime.datetime.now().strftime("%m/%d/%y")
	row_cells[2].text = 'Initial Revision'

	return document


def build_inputs_table(document,config):
	
	try:
		independantVars = config['variables']['independantVars']
	except:
		independantVars = config['independantVars']

	table = document.add_table(rows=1, cols=3)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Tag'
	hdr_cells[1].text = 'Normalization Range'
	hdr_cells[2].text = 'Valid Range'
	make_rows_bold(table.rows[0])

	for var in independantVars:
		tag = config['tag_normalize_dict'][str(var)]
		tagname = list(tag.keys())[0]
		normmax = tag[tagname][0]
		normmin = tag[tagname][1]
		try:
			min_conf = math.ceil(float(config["min_95th_pct"][str(var)]) * (normmax - normmin) + normmin)
			max_conf = math.floor(float(config["max_95th_pct"][str(var)]) * (normmax - normmin) + normmin)
		except:
			min_conf = math.ceil(float(config["min_training_range"][str(var)]) * (normmax - normmin) + normmin)
			max_conf = math.floor(float(config["max_training_range"][str(var)]) * (normmax - normmin) + normmin)

		row_cells = table.add_row().cells
		row_cells[0].text = tagname
		row_cells[1].text = str(normmin) +' - '+str(normmax)
		row_cells[2].text = str(min_conf)+ ' - ' + str(max_conf)

	return document


def build_outputs_table(document,config):
	
	try:
		var = config['variables']['dependantVar']
	except:
		var = config['dependantVar']

	output_table = document.add_table(rows=1, cols=3)
	hdr_cells = output_table.rows[0].cells
	hdr_cells[0].text = 'Tag'
	hdr_cells[1].text = 'Normalization Range'
	hdr_cells[2].text = 'Valid Range'
	make_rows_bold(output_table.rows[0])

	tag = config['tag_normalize_dict'][str(var)]
	tagname = list(tag.keys())[0]
	normmax = tag[tagname][0]
	normmin = tag[tagname][1]
	try:
		min_conf = math.ceil(float(config["min_training_range"][str(var)]) * (normmax - normmin) + normmin)
		max_conf = math.floor(float(config["max_training_range"][str(var)]) * (normmax - normmin) + normmin)
	except:
		min_conf = math.ceil(float(config["targetmin"]) * (normmax - normmin) + normmin)
		max_conf = math.floor(float(config["targetmax"]) * (normmax - normmin) + normmin)
		
	row_cells = output_table.add_row().cells
	row_cells[0].text = tagname
	row_cells[1].text = str(normmin) +' - '+str(normmax)
	row_cells[2].text = str(min_conf)+ ' - ' + str(max_conf)

	return document


def build_config_table(document,config):
	
	var = config['variables']['dependantVar']

	config_table = document.add_table(rows=6, cols=2)
	hdr_cells = config_table.columns[0].cells
	hdr_cells[0].text = 'Output Tag Name'
	hdr_cells[1].text = 'Training Scanrate'
	hdr_cells[2].text = 'Execution Scanrate'
	hdr_cells[3].text = 'Historical Depth'
	hdr_cells[4].text = 'Max Step'
	hdr_cells[5].text = 'Output Type'
	make_rows_bold(config_table.columns[0])

	tag = config['tag_normalize_dict'][str(var)]
	tagname = list(tag.keys())[0]
	trainingscanrate = config['training_scanrate']*config['data_sample_rate']
	executionscanrate = config['execution_scanrate']*config['data_sample_rate']

	col_cells = config_table.columns[1].cells
	col_cells[0].text = tagname
	col_cells[1].text = str(trainingscanrate) + ' seconds'
	col_cells[2].text = str(executionscanrate) + ' seconds'
	col_cells[3].text = str(config['agent_lookback']) + ' samples'
	col_cells[4].text = str(config['max_step']*100) + '%'
	col_cells[5].text = 'Positional'

	return document

def build_dt_config_table(document,config):
	
	config_table = document.add_table(rows=6, cols=2)
	hdr_cells = config_table.columns[0].cells
	hdr_cells[0].text = 'Output Tag Name'
	hdr_cells[1].text = 'Input Tags'
	hdr_cells[2].text = 'Training Scanrate'
	hdr_cells[3].text = 'Execution Scanrate'
	hdr_cells[4].text = 'Historical Depth'
	hdr_cells[5].text = 'Output Type'
	make_rows_bold(config_table.columns[0])

	dependantvar = config['dependantVar']
	independantVars = config['independantVars']
	dependanttag = config['tag_normalize_dict'][str(dependantvar)]

	independanttags = ''
	for var in independantVars:
		tag = config['tag_normalize_dict'][str(var)]
		tagname = list(tag.keys())[0]
		independanttags = independanttags + '\n' + tagname

	dependanttagname = list(dependanttag.keys())[0]
	scanrate = config['scanrate']*config['data_sample_rate']

	if config['velocity']:
		outputtype = 'Velocity'
	else:
		outputtype = 'Positional'

	col_cells = config_table.columns[1].cells
	col_cells[0].text = dependanttagname
	col_cells[1].text = independanttags[1:]
	col_cells[2].text = str(scanrate) + ' seconds'
	col_cells[3].text = str(scanrate) + ' seconds'
	col_cells[4].text = str(config['dt_lookback']) + ' samples'
	col_cells[5].text = outputtype

	return document

def build_csv_table(document,path_to_csv):

	csv_table = pd.read_csv(path_to_csv).astype(str)
	acr_table = document.add_table(rows=csv_table.shape[0], cols=csv_table.shape[1])
	hdr_cells = acr_table.columns[0].cells
	def_cells = acr_table.columns[1].cells
	for row in range(csv_table.shape[0]):
		hdr_cells[row].text = csv_table.iloc[row,0]
		def_cells[row].text = csv_table.iloc[row,1]
	make_rows_bold(acr_table.columns[0])

	return document